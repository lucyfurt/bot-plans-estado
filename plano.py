from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx2pdf import convert

# Carregar a planilha Excel
planilha_aulas = load_workbook('./plano.xlsx')
pagina_aulas = planilha_aulas['Sheet1']

# Iterar pelas linhas da planilha
for linha in pagina_aulas.iter_rows(min_row=2, values_only=True):
    try:
        periodo, aprendizagem, conteudos, procedimentos, avaliacao, etapa_ensino, ano, serie, disciplina, professor, *extras = linha
    except ValueError:
        print("Erro ao desempacotar linha:", linha)
        continue

    # Criar o documento do Word
    arquivo_word = Document()

    # Configurar a seção para ter um cabeçalho
    secao = arquivo_word.sections[0]
    header = secao.header
    header.is_linked_to_previous = False

    # Adicionar uma imagem centralizada no cabeçalho
    img_path = 'logo.png'
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run()
    picture = run.add_picture(img_path, width=Inches(8))

    # Criar estilo para o texto do plano
    estilo_texto = arquivo_word.styles['Normal']
    fonte = estilo_texto.font
    fonte.name = 'Arial'
    fonte.size = Pt(10)

    # Adicionar o texto ao documento com o estilo definido
    texto_plano = f"""
    ETAPA DE ENSINO: {etapa_ensino} ANO {ano}/SÉRIE: {serie} DISCIPLINA: {disciplina} PROFESSORA: {professor}
        """

    # Adicionar o texto ao documento com o estilo definido
    arquivo_word.add_paragraph(texto_plano, style='Normal')

    # Alterar a orientação do documento para paisagem
    secao.page_width = Inches(11)  # largura da página
    secao.page_height = Inches(8.5)  # altura da página

    # Adicionar tabela ao documento
    tabela = arquivo_word.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'  # Estilo da tabela

    # Definir cabeçalho da tabela
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = 'PERÍODO'
    cabecalho[1].text = 'APRENDIZAGENS BÁSICAS'
    cabecalho[2].text = 'CONTEÚDOS BÁSICOS'
    cabecalho[3].text = 'PROCEDIMENTOS METODOLÓGICOS'
    cabecalho[4].text = 'AVALIAÇÃO'

    # Adicionar dados à tabela a partir do Excel
    for linha_excel in pagina_aulas.iter_rows(min_row=2, values_only=True):
        dados_linha = linha_excel[:5]
        nova_linha = tabela.add_row().cells
        for i, dado in enumerate(dados_linha):
            nova_linha[i].text = str(dado)

    # Estilizar o cabeçalho da tabela
    for cell in cabecalho:
        cell.paragraphs[0].runs[0].font.bold = True 
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

    # Salvar o documento Word
   
    arquivo_word.save(f'./planos/plano_{disciplina}.docx')
   


